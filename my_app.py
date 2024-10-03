from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture  
document.add_picture(
                    'me.png.jpg',
                    width=Inches
                     
                     )

#name phone number and email details
name = input('What is your name : ')
speak('Hello ' + name + 'How are you today')
phone_Number = input('What is Phone Number : ')
email = input('What is your email Adderess : ')

document.add_paragraph(
    name + ' | ' + phone_Number + ' | '  + email)

#about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about Yourself?')
)

#work and experience
document.add_heading('working experience')

# more experiences
while True:
    has_more_experiences = input(
        'Do You Have more Experiences? YES or NO : ')
    if has_more_experiences.upper() == 'YES':
        p = document.add_paragraph()

        company = input('Enter company :')
        from_date = input('From Date :')
        to_date = input('TO Date :')

        p.add_run(company + '').bold =True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input (
        'Describe your experience at ' + company + ': '
        )
        p.add_run(experience_details)

    else:
        break



#Skills
document.add_heading= ('Working Skills')
skills =input('Enter Your Skill')
p= document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_Skills = input('Do you Have More Skills? YES or NO :')
    if has_more_Skills.upper() == 'YES':
        skill = input('Enter Skill')
        p =document.add_paragraph()
        p.style = 'List Bullet'

    else:
        break


#Footer
section = document.sections[0]
footer=section.footer
p = footer.paragraphs[0]
p.text = 'CV Generated using Amigosecode and Intuite Quickbooks course project'

document.save('cv.docx')